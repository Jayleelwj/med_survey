#!/usr/bin/env python3
"""Build the IEEE MS-MRI survey from one manuscript/content source.

The script deliberately keeps scientific content out of the renderer. Figure
captions, alt text, table cells, and equations come only from ``ieee_content``;
the manuscript supplies prose and numbered placement markers.  It produces a
template-based DOCX, an ``ieeecolor`` LaTeX source, the compiled PDF, and a
machine-readable build audit.
"""

from __future__ import annotations

import json
import math
import re
import shutil
import subprocess
import sys
import zipfile
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, Iterator, Sequence

from ieee_content import EQUATIONS, FIGURES, ROMAN_TABLE_NUMBERS, TABLES
from references_ieee import REFERENCES


ROOT = Path(__file__).resolve().parent
MANUSCRIPT = ROOT / "manuscript_ieee.md"
WORD_TEMPLATE = ROOT / "IEEE_word" / "Alternate_tj_template_ap.docx"
LATEX_TEMPLATE_DIR = ROOT / "IEEE_latex" / "alternate_tj_latex_template_ap"
FIGURE_DIR = ROOT / "output" / "figures_ieee"
OUTPUT_DIR = ROOT / "output"
LATEX_DIR = OUTPUT_DIR / "ieee_source"
LATEX_FIGURE_DIR = LATEX_DIR / "figures"
DOCX_OUT = OUTPUT_DIR / "Artificial_Intelligence_in_Multiple_Sclerosis_MRI_Survey_IEEE.docx"
PDF_OUT = OUTPUT_DIR / "Artificial_Intelligence_in_Multiple_Sclerosis_MRI_Survey_IEEE.pdf"
AUDIT_OUT = OUTPUT_DIR / "IEEE_BUILD_AUDIT.json"

PYTHON_RUNTIME = Path(
    "/Users/liwenjie/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3"
)
MATHML2OMML = Path("/Applications/Microsoft Word.app/Contents/Resources/mathml2omml.xsl")

TITLE_RE = re.compile(r"^#\s+(.+)$", re.MULTILINE)
CITATION_RE = re.compile(r"\[@([^\]]+)\]")
MARKER_RE = re.compile(r"^\[\[((?:FIGURE|TABLE|EQ)\d+)\]\]$")
ANY_MARKER_RE = re.compile(r"\[\[((?:FIGURE|TABLE|EQ)\d+)\)??\]\]")

# The deterministic vector generator retained descriptive stems for Figures
# 9--11. The aliases are explicit so the content source remains authoritative.
FIGURE_STEM_ALIASES = {
    "fig09_loss_error_map": "fig09_error_loss_mapping",
    "fig10_evaluation_plots": "fig10_evaluation_frameworks",
    "fig11_future_validation": "fig11_validation_roadmap",
}


def fail(message: str) -> "NoReturn":
    raise RuntimeError(message)


def normalize_space(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def strip_markdown_emphasis(text: str) -> str:
    return re.sub(r"\*\*(.*?)\*\*", r"\1", text).strip()


def extract_source() -> dict:
    if not MANUSCRIPT.exists():
        fail(f"Missing manuscript: {MANUSCRIPT}")
    source = MANUSCRIPT.read_text(encoding="utf-8")
    title_match = TITLE_RE.search(source)
    if not title_match:
        fail("Manuscript title was not found")

    abstract_match = re.search(
        r"^## Abstract\s*\n+(.*?)\n+\*\*Index Terms—\*\*(.*?)\n+## I\.",
        source,
        flags=re.MULTILINE | re.DOTALL,
    )
    if not abstract_match:
        fail("Abstract or Index Terms block was not found")

    body_start = source.find("## I.")
    declarations_start = source.find("## Declarations", body_start)
    if body_start < 0 or declarations_start < 0:
        fail("Body must run from '## I.' through '## Declarations'")
    body = source[body_start:].strip()
    if not body.startswith("## I.") or "## Declarations" not in body:
        fail("Body extraction boundaries are invalid")

    metadata = {}
    top = source[: source.find("## Abstract")]
    for label in ["Authors", "Affiliations", "ORCID", "Corresponding author", "Draft status"]:
        match = re.search(rf"^\*\*{re.escape(label)}:\*\*\s*(.+)$", top, re.MULTILINE)
        metadata[label] = normalize_space(match.group(1)) if match else "[NR]"

    abstract = normalize_space(abstract_match.group(1))
    index_terms = normalize_space(abstract_match.group(2)).strip(" .")
    return {
        "source": source,
        "title": normalize_space(title_match.group(1)),
        "metadata": metadata,
        "abstract": abstract,
        "index_terms": index_terms,
        "body": body,
        "blocks": parse_blocks(body),
    }


def parse_blocks(body: str) -> list[dict]:
    blocks: list[dict] = []
    for raw in re.split(r"\n\s*\n", body.strip()):
        block = normalize_space(raw)
        if not block:
            continue
        marker = MARKER_RE.fullmatch(block)
        if marker:
            blocks.append({"type": "marker", "key": marker.group(1), "text": block})
        elif block.startswith("### "):
            blocks.append({"type": "subsection", "text": block[4:].strip()})
        elif block.startswith("## "):
            blocks.append({"type": "section", "text": block[3:].strip()})
        else:
            blocks.append({"type": "paragraph", "text": block})
    return blocks


def validate_markers(source: dict) -> dict:
    text = source["body"]
    found = re.findall(r"\[\[((?:FIGURE|TABLE|EQ)\d+)\]\]", text)
    counts = Counter(found)
    expected = set(FIGURES) | set(TABLES) | set(EQUATIONS)
    actual = set(found)
    missing = sorted(expected - actual)
    unexpected = sorted(actual - expected)
    duplicates = sorted(key for key, count in counts.items() if count != 1)
    if missing or unexpected or duplicates:
        fail(
            "Placement-marker validation failed: "
            f"missing={missing}, unexpected={unexpected}, non_singletons={duplicates}"
        )
    return {
        "expected": len(expected),
        "found": len(found),
        "figures": len(FIGURES),
        "tables": len(TABLES),
        "equations": len(EQUATIONS),
        "all_unique_and_complete": True,
    }


def figure_number(key: str) -> int:
    return int(re.search(r"\d+", key).group())


def table_number(key: str) -> int:
    return int(re.search(r"\d+", key).group())


def equation_number(key: str) -> int:
    return int(re.search(r"\d+", key).group())


def resolve_figure_path(key: str, extension: str) -> Path:
    stem = FIGURES[key]["stem"]
    candidates = [stem]
    if stem in FIGURE_STEM_ALIASES:
        candidates.append(FIGURE_STEM_ALIASES[stem])

    manifest_path = FIGURE_DIR / "figure_manifest.json"
    if manifest_path.exists():
        try:
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            number = figure_number(key)
            for item in manifest.get("figures", []):
                if int(item.get("figure_number", -1)) == number:
                    candidates.append(str(item.get("stem", "")))
        except (ValueError, OSError, TypeError):
            pass

    seen = set()
    for candidate in candidates:
        if not candidate or candidate in seen:
            continue
        seen.add(candidate)
        path = FIGURE_DIR / f"{candidate}.{extension}"
        if path.exists():
            return path
    fail(f"Missing Figure {figure_number(key)} .{extension}; tried stems {sorted(seen)}")


def citation_keys(text: str) -> Iterator[str]:
    for group in CITATION_RE.findall(text):
        for item in group.split(";"):
            key = item.strip()
            if key.startswith("@"):
                key = key[1:]
            if key:
                yield key


def output_fragments(blocks: Sequence[dict]) -> Iterator[str]:
    for block in blocks:
        if block["type"] != "marker":
            yield block["text"]
            continue
        key = block["key"]
        if key in FIGURES:
            yield FIGURES[key]["caption"]
        elif key in TABLES:
            table = TABLES[key]
            yield table["caption"]
            yield from table["headers"]
            for row in table["rows"]:
                yield from row


def build_citation_order(blocks: Sequence[dict]) -> tuple[list[str], dict[str, int], list[str]]:
    order: list[str] = []
    seen = set()
    for fragment in output_fragments(blocks):
        for key in citation_keys(fragment):
            if key not in seen:
                seen.add(key)
                order.append(key)
    missing = sorted(key for key in order if key not in REFERENCES)
    if missing:
        fail(f"Undefined citation keys: {missing}")
    if len(order) < 100:
        fail(f"Only {len(order)} actually cited references; at least 100 are required")
    return order, {key: idx + 1 for idx, key in enumerate(order)}, missing


def latex_escape_plain(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
        "–": "--",
        "—": "---",
        "−": "-",
        "≤": r"$\leq$",
        "≥": r"$\geq$",
        "×": r"$\times$",
        "±": r"$\pm$",
        "¹": r"$^{1}$",
        "²": r"$^{2}$",
        "³": r"$^{3}$",
        "“": "``",
        "”": "''",
        "’": "'",
        "‘": "`",
        "→": r"$\rightarrow$",
        "↔": r"$\leftrightarrow$",
        "⊙": r"$\odot$",
        "ε": r"$\epsilon$",
        "α": r"$\alpha$",
        "β": r"$\beta$",
    }
    return "".join(replacements.get(char, char) for char in text)


LATEX_INLINE_RE = re.compile(
    r"(\*\*.*?\*\*|\$.*?\$|\[@[^\]]+\]|"
    r"\b(?:Figure|Fig\.)\s+\d+(?:\([a-z]\))?|\bTable\s+[IVX]+)"
)


def latex_inline(text: str) -> str:
    pieces: list[str] = []
    position = 0
    for match in LATEX_INLINE_RE.finditer(text):
        pieces.append(latex_escape_plain(text[position : match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            pieces.append(r"\textbf{" + latex_escape_plain(token[2:-2]) + "}")
        elif token.startswith("$"):
            pieces.append(token)
        elif token.startswith("[@"):
            keys = [item.strip().lstrip("@") for item in token[2:-1].split(";")]
            pieces.append(r"\cite{" + ",".join(keys) + "}")
        elif token.startswith("Table"):
            roman = token.split()[-1]
            inverse = {value: key for key, value in ROMAN_TABLE_NUMBERS.items()}
            pieces.append(r"Table~\ref{tab:" + str(inverse[roman]) + "}")
        else:
            number_match = re.search(r"\d+", token)
            number = int(number_match.group())
            suffix = token[number_match.end() :]
            pieces.append(r"Fig.~\ref{fig:" + str(number) + "}" + latex_escape_plain(suffix))
        position = match.end()
    pieces.append(latex_escape_plain(text[position:]))
    return "".join(pieces)


def latex_heading(text: str, level: str) -> str:
    if level == "section":
        if text == "Declarations":
            return r"\section*{Declarations}\label{sec:declarations}"
        clean = re.sub(r"^[IVX]+\.\s*", "", text)
        slug = re.sub(r"[^a-z0-9]+", "-", clean.lower()).strip("-")
        return rf"\section{{{latex_escape_plain(clean)}}}\label{{sec:{slug}}}"
    clean = re.sub(r"^[A-Z]\.\s*", "", text)
    slug = re.sub(r"[^a-z0-9]+", "-", clean.lower()).strip("-")
    return rf"\subsection{{{latex_escape_plain(clean)}}}\label{{subsec:{slug}}}"


def latex_figure(key: str) -> str:
    number = figure_number(key)
    pdf = resolve_figure_path(key, "pdf")
    destination = LATEX_FIGURE_DIR / pdf.name
    shutil.copy2(pdf, destination)
    caption = latex_inline(FIGURES[key]["caption"])
    max_height = "0.82\\textheight" if key == "FIGURE8" else "0.86\\textheight"
    # Figure 11 shares a float page with Table VI; the slightly narrower width
    # preserves the IEEE page box while retaining vector-label readability.
    max_width = "0.93\\textwidth" if key == "FIGURE11" else "\\textwidth"
    return "\n".join(
        [
            r"\begin{figure*}[!t]",
            r"\centering",
            rf"\includegraphics[width={max_width},height={max_height},keepaspectratio]{{figures/{destination.name}}}",
            rf"\caption{{{caption}}}",
            rf"\label{{fig:{number}}}",
            r"\end{figure*}",
        ]
    )


def latex_column_spec(column_count: int) -> str:
    if column_count <= 4:
        return "@{}" + "Y" * column_count + "@{}"
    return "@{}" + "Z" * column_count + "@{}"


def latex_table(key: str) -> str:
    number = table_number(key)
    table = TABLES[key]
    headers = table["headers"]
    rows = table["rows"]
    font = r"\scriptsize"
    output = [
        r"\begin{table*}[!t]",
        r"\centering",
        rf"\caption{{{latex_escape_plain(table['caption'])}}}",
        rf"\label{{tab:{number}}}",
        font,
        r"\setlength{\tabcolsep}{2.2pt}",
        r"\renewcommand{\arraystretch}{1.10}",
        rf"\begin{{tabularx}}{{\textwidth}}{{{latex_column_spec(len(headers))}}}",
        r"\toprule",
        " & ".join(r"\textbf{" + latex_inline(cell) + "}" for cell in headers) + r" \\",
        r"\midrule",
    ]
    for row in rows:
        if len(row) != len(headers):
            fail(f"{key} row has {len(row)} cells; expected {len(headers)}")
        output.append(" & ".join(latex_inline(cell) for cell in row) + r" \\")
    output.extend([r"\bottomrule", r"\end{tabularx}", r"\end{table*}"])
    return "\n".join(output)


def latex_equation(key: str) -> str:
    equation = EQUATIONS[key]
    return "\n".join(
        [
            r"\begin{equation}",
            equation["latex"],
            rf"\label{{{equation['label']}}}",
            r"\end{equation}",
        ]
    )


def latex_body(blocks: Sequence[dict]) -> str:
    rendered: list[str] = []
    index = 0
    while index < len(blocks):
        block = blocks[index]
        kind = block["type"]
        if kind == "section":
            # Flush the preceding top-level section's starred floats, but let
            # prose flow around figures and tables within a section. Per-marker
            # barriers create sparsely filled pages in two-column IEEE layouts.
            if rendered and re.match(r"^(?:III|V)\.\s", block["text"]):
                rendered.append(r"\FloatBarrier")
                if re.match(r"^V\.\s", block["text"]):
                    rendered.append(r"\clearpage")
            rendered.append(latex_heading(block["text"], "section"))
        elif kind == "subsection":
            rendered.append(latex_heading(block["text"], "subsection"))
        elif kind == "paragraph":
            rendered.append(latex_inline(block["text"]) + "\n")
        elif block["key"] in FIGURES or block["key"] in TABLES:
            # Starred floats use separate figure and table queues in two-column
            # mode. Keep consecutive markers together; top-level section
            # barriers below prevent them from crossing major section bounds.
            while index < len(blocks):
                marker = blocks[index]
                if marker["type"] != "marker" or not (
                    marker["key"] in FIGURES or marker["key"] in TABLES
                ):
                    break
                key = marker["key"]
                rendered.append(latex_figure(key) if key in FIGURES else latex_table(key))
                index += 1
            continue
        else:
            key = block["key"]
            if key in EQUATIONS:
                rendered.append(latex_equation(key))
            else:
                fail(f"Unknown marker {key}")
        index += 1
    return "\n\n".join(rendered)


def latex_reference(reference: str) -> str:
    escaped = latex_escape_plain(reference)
    # Permit safe breaks in long DOI and URL-like tokens without changing text.
    escaped = escaped.replace("doi:", r"doi:\allowbreak ")
    escaped = escaped.replace("https://", r"https://\allowbreak ")
    escaped = escaped.replace("/", r"/\allowbreak ")
    return escaped


def latex_author_text(text: str) -> str:
    protected = text.replace("¹", "<<SUP1>>").replace("²", "<<SUP2>>").replace("³", "<<SUP3>>")
    escaped = latex_escape_plain(protected)
    return (
        escaped.replace("<<SUP1>>", r"\textsuperscript{1}")
        .replace("<<SUP2>>", r"\textsuperscript{2}")
        .replace("<<SUP3>>", r"\textsuperscript{3}")
    )


def plain_pdf_author(text: str) -> str:
    return strip_markdown_emphasis(text).translate(str.maketrans({"¹": "1", "²": "2", "³": "3"}))


def write_blank_logo() -> None:
    # ieeecolor requests the logo file even when the configured width is zero.
    (LATEX_DIR / "blanklogo.eps").write_text(
        "%!PS-Adobe-3.0 EPSF-3.0\n"
        "%%BoundingBox: 0 0 1 1\n"
        "%%HiResBoundingBox: 0 0 1 1\n"
        "newpath\nshowpage\n%%EOF\n",
        encoding="ascii",
    )


def prepare_latex_template() -> None:
    LATEX_DIR.mkdir(parents=True, exist_ok=True)
    LATEX_FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(LATEX_TEMPLATE_DIR / "ieeecolor.cls", LATEX_DIR / "ieeecolor.cls")
    generic = (LATEX_TEMPLATE_DIR / "generic.sty").read_text(encoding="utf-8")
    generic = re.sub(r"\\def\\logoname\{[^}]*\}", r"\\def\\logoname{blanklogo}", generic)
    generic = re.sub(
        r"\\def\\journalname\{[^}]*\}",
        r"\\def\\journalname{AI-ASSISTED RESEARCH DRAFT}",
        generic,
    )
    (LATEX_DIR / "generic.sty").write_text(generic, encoding="utf-8")
    write_blank_logo()


def render_latex(source: dict, citation_order: Sequence[str]) -> Path:
    prepare_latex_template()
    metadata = source["metadata"]
    references = "\n".join(
        rf"\bibitem{{{key}}} {latex_reference(REFERENCES[key])}" for key in citation_order
    )
    index_terms = source["index_terms"].replace(";", ",")
    author_text = latex_author_text(metadata["Authors"])
    thanks = [
        metadata["Affiliations"],
        "ORCID: " + metadata["ORCID"],
        "Corresponding author: " + metadata["Corresponding author"],
        metadata["Draft status"],
    ]
    thanks_text = "\n".join(rf"\thanks{{{latex_escape_plain(item)}}}" for item in thanks)
    document = rf"""\documentclass[journal,twoside,web]{{ieeecolor}}
\usepackage{{generic}}
\usepackage{{cite}}
\usepackage{{amsmath,amssymb,amsfonts}}
\usepackage{{graphicx}}
\usepackage{{booktabs,tabularx,array}}
\usepackage{{placeins}}
\usepackage[T1]{{fontenc}}
\usepackage[utf8]{{inputenc}}
\usepackage{{microtype}}
\usepackage{{url}}
\usepackage{{hyperref}}
\pdfstringdefDisableCommands{{\def\thanks#1{{}}\def\textsuperscript#1{{#1}}}}
\hypersetup{{
  hidelinks,
  pdftitle={{{latex_escape_plain(source['title'])}}},
  pdfauthor={{{latex_escape_plain(plain_pdf_author(metadata['Authors']))}}},
  pdfsubject={{AI-assisted survey draft on artificial intelligence in multiple sclerosis MRI}},
  pdfkeywords={{{latex_escape_plain(index_terms)}}}
}}
\newcolumntype{{Y}}{{>{{\raggedright\arraybackslash\hspace{{0pt}}}}X}}
\newcolumntype{{Z}}{{>{{\raggedright\arraybackslash\hspace{{0pt}}}}X}}
\providecommand{{\refname}}{{References}}
\emergencystretch=2em
\Urlmuskip=0mu plus 1mu
\makeatletter
\def\ps@headings{{%
  \def\@oddhead{{\scriptsize\rightmark\hfil\thepage}}%
  \def\@evenhead{{\scriptsize\thepage\hfil\leftmark}}%
  \def\@oddfoot{{}}\def\@evenfoot{{}}}}
\def\ps@titlepagestyle{{%
  \def\@oddhead{{\scriptsize AI-ASSISTED RESEARCH DRAFT\hfil\thepage}}%
  \def\@evenhead{{\scriptsize\thepage\hfil AI-ASSISTED RESEARCH DRAFT}}%
  \def\@oddfoot{{}}\def\@evenfoot{{}}}}
\makeatother
\markboth{{AI-ASSISTED RESEARCH DRAFT}}{{Author et al.: MS MRI AI Survey}}

\begin{{document}}
\title{{{latex_escape_plain(source['title'])}}}
\author{{{author_text}
{thanks_text}}}
\maketitle
\pagestyle{{headings}}

\begin{{abstract}}
{latex_inline(source['abstract'])}
\end{{abstract}}

\begin{{IEEEkeywords}}
{latex_escape_plain(index_terms)}
\end{{IEEEkeywords}}

{latex_body(source['blocks'])}

\FloatBarrier
\clearpage
\begin{{thebibliography}}{{999}}
{references}
\end{{thebibliography}}

\end{{document}}
"""
    tex_path = LATEX_DIR / "main.tex"
    tex_path.write_text(document, encoding="utf-8")
    return tex_path


def run_command(command: Sequence[str], cwd: Path | None = None) -> dict:
    result = subprocess.run(
        list(command),
        cwd=str(cwd) if cwd else None,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        check=False,
    )
    return {
        "command": list(command),
        "returncode": result.returncode,
        "output": result.stdout,
    }


def compile_latex(tex_path: Path) -> tuple[Path, dict]:
    result = run_command(
        [
            "latexmk",
            "-pdf",
            "-halt-on-error",
            "-file-line-error",
            "-interaction=nonstopmode",
            tex_path.name,
        ],
        cwd=LATEX_DIR,
    )
    log_path = LATEX_DIR / "main.log"
    log = log_path.read_text(encoding="utf-8", errors="ignore") if log_path.exists() else ""
    compile_audit = {
        "returncode": result["returncode"],
        "undefined_references": len(re.findall(r"undefined references?", log, re.IGNORECASE)),
        "undefined_citations": len(re.findall(r"Citation .* undefined", log, re.IGNORECASE)),
        "missing_files": len(re.findall(r"File `[^']+' not found", log)),
        "overfull_hbox": len(re.findall(r"Overfull \\hbox", log)),
        "overfull_vbox": len(re.findall(r"Overfull \\vbox", log)),
        "underfull_hbox": len(re.findall(r"Underfull \\hbox", log)),
        "underfull_vbox": len(re.findall(r"Underfull \\vbox", log)),
        "tail": result["output"][-8000:],
    }
    built_pdf = LATEX_DIR / "main.pdf"
    if result["returncode"] != 0 or not built_pdf.exists():
        fail("LaTeX compilation failed:\n" + result["output"][-6000:])
    shutil.copy2(built_pdf, PDF_OUT)
    return built_pdf, compile_audit


def import_docx_dependencies():
    try:
        from docx import Document  # noqa: F401
        from lxml import etree  # noqa: F401
    except ImportError as exc:
        fail(
            "python-docx/lxml are required. Run with the bundled workspace Python: "
            f"{PYTHON_RUNTIME} {Path(__file__).name}"
        )


def clear_document_body(document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_section_columns(section, count: int, gap_twips: int = 720) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(gap_twips if count > 1 else 0))
    cols.set(qn("w:equalWidth"), "1")


def apply_page_geometry(section) -> None:
    from docx.shared import Inches

    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.top_margin = Inches(0.70)
    section.bottom_margin = Inches(0.70)
    section.header_distance = Inches(0.30)
    section.footer_distance = Inches(0.30)
    section.gutter = Inches(0)


def clear_story(story) -> None:
    element = story._element
    for child in list(element):
        element.remove(child)


def add_simple_field(paragraph, instruction: str, fallback: str, *, bold: bool = False):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    field.set(qn("w:dirty"), "true")
    run = OxmlElement("w:r")
    if bold:
        r_pr = OxmlElement("w:rPr")
        bold_element = OxmlElement("w:b")
        r_pr.append(bold_element)
        run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = fallback
    run.append(text)
    field.append(run)
    paragraph._p.append(field)
    return field


def add_bookmarked_run(paragraph, text: str, name: str, bookmark_id: int, *, bold: bool = False):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    paragraph._p.append(start)
    run = paragraph.add_run(text)
    run.bold = bold
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(end)
    return run


def inline_math_to_unicode(text: str) -> str:
    mapping = {
        r"y=F(x)+x": "y = F(x) + x",
        r"L_1=\frac{1}{N}\sum_i|x_i-\hat{x}_i|": "L₁ = (1/N) Σᵢ |xᵢ − x̂ᵢ|",
        r"L_2=\frac{1}{N}\sum_i(x_i-\hat{x}_i)^2": "L₂ = (1/N) Σᵢ (xᵢ − x̂ᵢ)²",
        r"y_i\in\{0,1\}": "yᵢ ∈ {0, 1}",
        r"p_i": "pᵢ",
        r"\epsilon": "ε",
        r"(1-p_t)^\gamma": "(1 − pₜ)ᵞ",
    }
    return mapping.get(text, text.replace("\\", ""))


WORD_INLINE_RE = re.compile(
    r"(\*\*.*?\*\*|\$.*?\$|\[@[^\]]+\]|"
    r"\b(?:Figure|Fig\.)\s+\d+(?:\([a-z]\))?|\bTable\s+[IVX]+)"
)


def add_word_inline(paragraph, text: str, citation_numbers: dict[str, int], *, base_bold: bool = False) -> None:
    position = 0
    inverse_roman = {value: key for key, value in ROMAN_TABLE_NUMBERS.items()}
    for match in WORD_INLINE_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            run.bold = base_bold
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("$"):
            run = paragraph.add_run(inline_math_to_unicode(token[1:-1]))
            run.italic = True
            run.bold = base_bold
        elif token.startswith("[@"):
            keys = [item.strip().lstrip("@") for item in token[2:-1].split(";")]
            numbers = [citation_numbers[key] for key in keys]
            run = paragraph.add_run("[" + "], [".join(str(number) for number in numbers) + "]")
            run.bold = base_bold
        elif token.startswith("Table"):
            roman = token.split()[-1]
            number = inverse_roman[roman]
            run = paragraph.add_run("Table ")
            run.bold = base_bold
            add_simple_field(paragraph, f" REF tab_{number} \\h ", roman, bold=base_bold)
        else:
            number_match = re.search(r"\d+", token)
            number = int(number_match.group())
            suffix = token[number_match.end() :]
            prefix = "Fig. " if token.startswith("Fig.") else "Figure "
            run = paragraph.add_run(prefix)
            run.bold = base_bold
            add_simple_field(paragraph, f" REF fig_{number} \\h ", str(number), bold=base_bold)
            if suffix:
                run = paragraph.add_run(suffix)
                run.bold = base_bold
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        run.bold = base_bold


def add_omml_equation(document, key: str):
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Inches
    from lxml import etree

    if not MATHML2OMML.exists():
        fail(f"Missing Word MathML-to-OMML transform: {MATHML2OMML}")
    transform = etree.XSLT(etree.parse(str(MATHML2OMML)))
    mathml = etree.fromstring(EQUATIONS[key]["mathml"].encode("utf-8"))
    omml = transform(mathml).getroot()

    paragraph = document.add_paragraph(style="Equation")
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(3.58), WD_TAB_ALIGNMENT.CENTER)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.05), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.add_run("\t")
    paragraph._p.append(omml)
    paragraph.add_run(f"\t({equation_number(key)})")
    return paragraph


def image_dimensions(path: Path, target_width_inches: float = 7.16) -> tuple[float, float]:
    from PIL import Image

    with Image.open(path) as image:
        width, height = image.size
    return target_width_inches, target_width_inches * height / width


def add_figure(document, key: str, bookmark_id: int):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    number = figure_number(key)
    png = resolve_figure_path(key, "png")
    width, height = image_dimensions(png)
    if height > 8.70:
        scale = 8.70 / height
        width *= scale
        height *= scale

    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.keep_with_next = True
    image_run = image_paragraph.add_run()
    inline = image_run.add_picture(str(png), width=Inches(width), height=Inches(height))
    inline._inline.docPr.set("descr", FIGURES[key]["alt"])
    inline._inline.docPr.set("title", f"Figure {number}")

    caption = document.add_paragraph(style="FigureCaption")
    caption.paragraph_format.keep_together = True
    caption.add_run("Fig. ")
    add_bookmarked_run(caption, str(number), f"fig_{number}", bookmark_id, bold=True)
    caption.add_run(". ")
    add_word_inline(caption, FIGURES[key]["caption"], {})
    return bookmark_id + 1


def table_widths(table: dict, total_twips: int = 10310) -> list[int]:
    columns = len(table["headers"])
    weights = []
    for column in range(columns):
        lengths = [len(str(table["headers"][column]))]
        lengths.extend(len(str(row[column])) for row in table["rows"])
        representative = max(12.0, 0.55 * max(lengths) + 0.45 * sum(lengths) / len(lengths))
        weights.append(math.sqrt(representative))
    minimum = 720 if columns >= 7 else 1000
    free = total_twips - minimum * columns
    widths = [minimum + int(free * weight / sum(weights)) for weight in weights]
    widths[-1] += total_twips - sum(widths)
    return widths


def set_cell_border(cell, **edges) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, settings in edges.items():
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        for attribute, value in settings.items():
            element.set(qn(f"w:{attribute}"), str(value))


def set_table_borders(table) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table_pr = table._tbl.tblPr
    borders = table_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_pr.append(borders)
    for edge in ["left", "right", "insideH", "insideV"]:
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    for edge in ["top", "bottom"]:
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:color"), "000000")
        borders.append(element)


def add_table(document, key: str, citation_numbers: dict[str, int], bookmark_id: int):
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, Twips

    number = table_number(key)
    roman = ROMAN_TABLE_NUMBERS[number]
    content = TABLES[key]

    title = document.add_paragraph(style="TableTitle")
    title.paragraph_format.keep_with_next = True
    title.add_run("TABLE ")
    add_bookmarked_run(title, roman, f"tab_{number}", bookmark_id, bold=True)
    caption = document.add_paragraph(style="TableTitle")
    caption.paragraph_format.keep_with_next = True
    caption.add_run(content["caption"]).bold = True

    rows = [content["headers"], *content["rows"]]
    widths = table_widths(content)
    word_table = document.add_table(rows=len(rows), cols=len(content["headers"]))
    word_table.autofit = False
    word_table.alignment = 1
    set_table_borders(word_table)

    table_pr = word_table._tbl.tblPr
    table_width = OxmlElement("w:tblW")
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    table_pr.append(table_width)

    for row_index, row_values in enumerate(rows):
        row = word_table.rows[row_index]
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        if row_index == 0:
            row._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for column, value in enumerate(row_values):
            cell = row.cells[column]
            cell.width = Twips(widths[column])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.find(qn("w:tcW"))
            tc_width.set(qn("w:w"), str(widths[column]))
            tc_width.set(qn("w:type"), "dxa")
            margins = OxmlElement("w:tcMar")
            for side, amount in [("top", 50), ("bottom", 50), ("left", 55), ("right", 55)]:
                margin = OxmlElement(f"w:{side}")
                margin.set(qn("w:w"), str(amount))
                margin.set(qn("w:type"), "dxa")
                margins.append(margin)
            tc_pr.append(margins)
            if row_index == 0:
                set_cell_border(cell, bottom={"val": "single", "sz": "6", "color": "000000"})

            paragraph = cell.paragraphs[0]
            paragraph.style = "Text"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            add_word_inline(paragraph, str(value), citation_numbers, base_bold=row_index == 0)
            for run in paragraph.runs:
                run.font.size = Pt(6.5 if len(content["headers"]) >= 7 else 7.0)
    return bookmark_id + 1


def prepare_word_output_styles(document) -> None:
    from docx.oxml.ns import qn

    # The supplied IEEE template positions its sample Title and Authors styles
    # in text frames. Remove only those layout properties so generated title
    # matter remains template-styled but participates in normal document flow.
    for style_name in ("Title", "Authors"):
        style_ppr = document.styles[style_name]._element.get_or_add_pPr()
        frame = style_ppr.find(qn("w:framePr"))
        if frame is not None:
            style_ppr.remove(frame)

    # The template's References style carries sample automatic numbering. The
    # renderer supplies IEEE bracket numbers from the citation ledger, so the
    # inherited numPr would duplicate numbers in Word.
    references_ppr = document.styles["References"]._element.get_or_add_pPr()
    numbering = references_ppr.find(qn("w:numPr"))
    if numbering is not None:
        references_ppr.remove(numbering)


def add_title_matter(document, source: dict) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 1.05
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run(source["title"])
    title_run.font.size = Pt(20)

    authors = document.add_paragraph(style="Authors")
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.paragraph_format.space_after = Pt(3)
    author_run = authors.add_run(source["metadata"]["Authors"])
    author_run.font.size = Pt(10)
    for label in ["Affiliations", "ORCID", "Corresponding author"]:
        paragraph = document.add_paragraph(style="Authors")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(1.5)
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        label_run.font.size = Pt(9)
        value_run = paragraph.add_run(source["metadata"][label])
        value_run.font.size = Pt(9)

    status = document.add_paragraph(style="Authors")
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status.paragraph_format.space_after = Pt(6)
    run = status.add_run(source["metadata"]["Draft status"])
    run.italic = True
    run.font.size = Pt(9)

    abstract = document.add_paragraph(style="Abstract")
    abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract.add_run("Abstract—").bold = True
    content_run = abstract.add_run(source["abstract"])
    content_run.bold = False

    terms = document.add_paragraph(style="IndexTerms")
    terms.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    terms.add_run("Index Terms—").bold = True
    terms_run = terms.add_run(source["index_terms"])
    terms_run.bold = False


def add_body_paragraph(document, block: dict, citation_numbers: dict[str, int]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if block["type"] == "section":
        paragraph = document.add_paragraph(style="Heading1")
        text = block["text"]
        if text == "Declarations":
            paragraph.add_run("DECLARATIONS")
        else:
            match = re.match(r"^([IVX]+)\.\s*(.*)$", text)
            paragraph.add_run(f"{match.group(1)}. {match.group(2).upper()}" if match else text.upper())
    elif block["type"] == "subsection":
        paragraph = document.add_paragraph(style="Heading2")
        paragraph.add_run(block["text"])
    else:
        paragraph = document.add_paragraph(style="Text")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_word_inline(paragraph, block["text"], citation_numbers)


def add_page_number_footer(document) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    first = document.sections[0]
    clear_story(first.header)
    clear_story(first.footer)
    footer_paragraph = first.footer.add_paragraph()
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_simple_field(footer_paragraph, " PAGE ", "1")
    for index, section in enumerate(document.sections):
        page_numbering = section._sectPr.find(qn("w:pgNumType"))
        if page_numbering is None:
            page_numbering = OxmlElement("w:pgNumType")
            section._sectPr.append(page_numbering)
        if index == 0:
            page_numbering.set(qn("w:start"), "1")
        else:
            page_numbering.attrib.pop(qn("w:start"), None)
    for section in document.sections[1:]:
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True


def set_update_fields(document) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_word_references(document, citation_order: Sequence[str]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    heading = document.add_paragraph(style="ReferenceHead")
    heading.paragraph_format.keep_with_next = True
    heading.add_run("REFERENCES")
    for index, key in enumerate(citation_order, 1):
        paragraph = document.add_paragraph(style="References")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.first_line_indent = Inches(-0.18)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 0.95
        paragraph.paragraph_format.keep_together = False
        paragraph.paragraph_format.keep_with_next = False
        paragraph.paragraph_format.widow_control = False
        run = paragraph.add_run(f"[{index}] {REFERENCES[key]}")
        run.font.size = Pt(7.5)


def render_docx(source: dict, citation_order: Sequence[str], citation_numbers: dict[str, int]) -> dict:
    import_docx_dependencies()
    from docx import Document
    from docx.enum.section import WD_SECTION

    if not WORD_TEMPLATE.exists():
        fail(f"Missing Word template: {WORD_TEMPLATE}")
    document = Document(str(WORD_TEMPLATE))
    clear_document_body(document)
    prepare_word_output_styles(document)
    first = document.sections[0]
    apply_page_geometry(first)
    set_section_columns(first, 1)
    add_title_matter(document, source)

    body_section = document.add_section(WD_SECTION.CONTINUOUS)
    apply_page_geometry(body_section)
    set_section_columns(body_section, 2, 720)

    bookmark_id = 1
    blocks = source["blocks"]
    index = 0
    while index < len(blocks):
        block = blocks[index]
        is_full_width = block["type"] == "marker" and (
            block["key"] in FIGURES or block["key"] in TABLES
        )
        if not is_full_width:
            if block["type"] == "marker" and block["key"] in EQUATIONS:
                add_omml_equation(document, block["key"])
            else:
                add_body_paragraph(document, block, citation_numbers)
            index += 1
            continue

        one_column = document.add_section(WD_SECTION.CONTINUOUS)
        apply_page_geometry(one_column)
        set_section_columns(one_column, 1)
        while index < len(blocks):
            marker = blocks[index]
            if marker["type"] != "marker" or not (
                marker["key"] in FIGURES or marker["key"] in TABLES
            ):
                break
            if marker["key"] in FIGURES:
                bookmark_id = add_figure(document, marker["key"], bookmark_id)
            else:
                if marker["key"] == "TABLE2":
                    document.add_page_break()
                bookmark_id = add_table(
                    document, marker["key"], citation_numbers, bookmark_id
                )
            index += 1
        if index < len(blocks):
            two_column = document.add_section(WD_SECTION.CONTINUOUS)
            apply_page_geometry(two_column)
            set_section_columns(two_column, 2, 720)

    add_word_references(document, citation_order)
    add_page_number_footer(document)
    set_update_fields(document)
    document.settings.odd_and_even_pages_header_footer = False
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(DOCX_OUT))
    return inspect_docx(DOCX_OUT)


def inspect_docx(path: Path) -> dict:
    import_docx_dependencies()
    from lxml import etree

    result = {
        "path": str(path),
        "size_bytes": path.stat().st_size,
        "zip_integrity": False,
        "xml_well_formed": False,
    }
    with zipfile.ZipFile(path) as archive:
        result["zip_integrity"] = archive.testzip() is None
        xml_names = [name for name in archive.namelist() if name.endswith(".xml")]
        for name in xml_names:
            etree.fromstring(archive.read(name))
        result["xml_well_formed"] = True
        document_xml = etree.fromstring(archive.read("word/document.xml"))
        styles_xml = etree.fromstring(archive.read("word/styles.xml"))
        ns = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
            "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        }
        page_fields = 0
        for footer_name in [name for name in archive.namelist() if re.fullmatch(r"word/footer\d+\.xml", name)]:
            footer_xml = etree.fromstring(archive.read(footer_name))
            page_fields += len(
                footer_xml.xpath('.//w:fldSimple[contains(@w:instr,"PAGE")]', namespaces=ns)
            )
        result.update(
            {
                "sections": len(document_xml.xpath(".//w:sectPr", namespaces=ns)),
                "two_column_sections": len(
                    document_xml.xpath('.//w:cols[@w:num="2"]', namespaces=ns)
                ),
                "single_column_sections": len(
                    document_xml.xpath('.//w:cols[@w:num="1"]', namespaces=ns)
                ),
                "column_gap_720_twips": len(
                    document_xml.xpath('.//w:cols[@w:num="2"][@w:space="720"]', namespaces=ns)
                ),
                "native_omml_equations": len(document_xml.xpath(".//m:oMath", namespaces=ns)),
                "bookmarks": len(document_xml.xpath(".//w:bookmarkStart", namespaces=ns)),
                "ref_fields": len(
                    document_xml.xpath('.//w:fldSimple[contains(@w:instr,"REF")]', namespaces=ns)
                ),
                "page_fields": page_fields,
                "page_number_restarts": len(
                    document_xml.xpath('.//w:pgNumType[@w:start]', namespaces=ns)
                ),
                "table_rows": len(document_xml.xpath('.//w:tbl/w:tr', namespaces=ns)),
                "non_splittable_table_rows": len(
                    document_xml.xpath('.//w:tbl/w:tr[w:trPr/w:cantSplit]', namespaces=ns)
                ),
                "figures_with_alt_text": len(
                    document_xml.xpath('.//wp:docPr[@descr!=""]', namespaces=ns)
                ),
                "floating_frames_in_body": len(
                    document_xml.xpath(".//w:framePr", namespaces=ns)
                ),
                "title_author_style_frames": len(
                    styles_xml.xpath(
                        './/w:style[@w:styleId="Title" or @w:styleId="Authors"]//w:framePr',
                        namespaces=ns,
                    )
                ),
                "references_style_numbering": len(
                    styles_xml.xpath(
                        './/w:style[@w:styleId="References"]//w:numPr', namespaces=ns
                    )
                ),
            }
        )
    return result


def inspect_pdf(path: Path) -> dict:
    audit = {"path": str(path), "size_bytes": path.stat().st_size}
    for tool, arguments, key in [
        ("pdfinfo", [str(path)], "pdfinfo"),
        ("pdffonts", [str(path)], "pdffonts"),
        ("pdfimages", ["-list", str(path)], "pdfimages"),
    ]:
        result = run_command([tool, *arguments])
        audit[key] = {
            "returncode": result["returncode"],
            "output": result["output"],
        }
    page_match = re.search(r"^Pages:\s+(\d+)", audit["pdfinfo"]["output"], re.MULTILINE)
    audit["pages"] = int(page_match.group(1)) if page_match else None
    audit["fonts_embedded_or_subset"] = all(
        "yes" in line.lower()
        for line in audit["pdffonts"]["output"].splitlines()[2:]
        if line.strip()
    )
    return audit


def main() -> int:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    source = extract_source()
    marker_audit = validate_markers(source)
    citation_order, citation_numbers, missing_citations = build_citation_order(source["blocks"])

    figure_audit = {}
    for key in FIGURES:
        pdf = resolve_figure_path(key, "pdf")
        png = resolve_figure_path(key, "png")
        figure_audit[key] = {
            "number": figure_number(key),
            "content_stem": FIGURES[key]["stem"],
            "resolved_pdf": str(pdf),
            "resolved_png": str(png),
            "pdf_bytes": pdf.stat().st_size,
            "png_bytes": png.stat().st_size,
        }

    tex_path = render_latex(source, citation_order)
    built_pdf, latex_compile = compile_latex(tex_path)
    docx_audit = render_docx(source, citation_order, citation_numbers)
    pdf_audit = inspect_pdf(PDF_OUT)

    abstract_words = len(re.findall(r"\b[\w'-]+\b", source["abstract"]))
    cited_set = set(citation_order)
    audit = {
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "status": "pass",
        "source": {
            "manuscript": str(MANUSCRIPT),
            "content_module": str(ROOT / "ieee_content.py"),
            "reference_module": str(ROOT / "references_ieee.py"),
            "word_template": str(WORD_TEMPLATE),
            "latex_template": str(LATEX_TEMPLATE_DIR),
            "abstract_words": abstract_words,
        },
        "markers": marker_audit,
        "citations": {
            "actually_cited": len(citation_order),
            "minimum_met": len(citation_order) >= 100,
            "undefined": missing_citations,
            "order": citation_order,
            "catalog_entries_omitted_as_uncited": sorted(set(REFERENCES) - cited_set),
        },
        "figures": figure_audit,
        "latex": {
            "source": str(tex_path),
            "compiled_pdf": str(built_pdf),
            "template_class": "ieeecolor",
            "thebibliography_width": 999,
            "visible_demo_logo": False,
            "compile": latex_compile,
        },
        "word": docx_audit,
        "pdf": pdf_audit,
        "acceptance": {
            "all_markers_complete": marker_audit["all_unique_and_complete"],
            "at_least_100_cited_references": len(citation_order) >= 100,
            "latex_compiled": latex_compile["returncode"] == 0,
            "no_undefined_references": latex_compile["undefined_references"] == 0,
            "no_undefined_citations": latex_compile["undefined_citations"] == 0,
            "no_overfull_boxes": (
                latex_compile["overfull_hbox"] == 0
                and latex_compile["overfull_vbox"] == 0
            ),
            "docx_zip_integrity": docx_audit["zip_integrity"],
            "docx_xml_well_formed": docx_audit["xml_well_formed"],
            "three_native_omml_equations": docx_audit["native_omml_equations"] == 3,
            "all_figures_have_alt_text": docx_audit["figures_with_alt_text"] == len(FIGURES),
            "word_ref_fields_present": docx_audit["ref_fields"] > 0,
            "word_page_field_present": docx_audit["page_fields"] > 0,
            "word_page_numbers_continuous": docx_audit["page_number_restarts"] == 1,
            "word_table_rows_do_not_split": (
                docx_audit["table_rows"] == docx_audit["non_splittable_table_rows"]
            ),
            "half_inch_column_gap": docx_audit["column_gap_720_twips"] > 0,
            "word_title_matter_in_normal_flow": (
                docx_audit["floating_frames_in_body"] == 0
                and docx_audit["title_author_style_frames"] == 0
            ),
            "word_reference_numbers_not_duplicated": (
                docx_audit["references_style_numbering"] == 0
            ),
        },
    }
    audit["status"] = "pass" if all(audit["acceptance"].values()) else "fail"
    AUDIT_OUT.write_text(json.dumps(audit, indent=2, ensure_ascii=False), encoding="utf-8")

    print(f"DOCX: {DOCX_OUT}")
    print(f"PDF:  {PDF_OUT}")
    print(f"TeX:  {tex_path}")
    print(f"Audit: {AUDIT_OUT}")
    print(
        f"Citations: {len(citation_order)}; figures: {len(FIGURES)}; "
        f"tables: {len(TABLES)}; equations: {len(EQUATIONS)}; pages: {pdf_audit['pages']}"
    )
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"BUILD FAILED: {exc}", file=sys.stderr)
        raise
