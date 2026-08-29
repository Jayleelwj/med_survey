import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT


root = Path(__file__).resolve().parents[1]
docx_path = root / "output" / "Artificial_Intelligence_in_Multiple_Sclerosis_MRI_Survey.docx"
pdf_path = root / "output" / "Artificial_Intelligence_in_Multiple_Sclerosis_MRI_Survey.pdf"
audit_path = root / "output" / "manuscript_audit.json"

document = Document(docx_path)
audit = json.loads(audit_path.read_text(encoding="utf-8"))
paragraph_text = "\n".join(p.text for p in document.paragraphs)

with zipfile.ZipFile(docx_path) as archive:
    document_xml = archive.read("word/document.xml").decode("utf-8")
    footer_xml = "\n".join(
        archive.read(name).decode("utf-8")
        for name in archive.namelist()
        if re.fullmatch(r"word/footer\d+\.xml", name)
    )

checks = {
    "docx_exists": docx_path.exists() and docx_path.stat().st_size > 0,
    "pdf_exists": pdf_path.exists() and pdf_path.stat().st_size > 0,
    "title_present": "A Survey of Artificial Intelligence in Multiple Sclerosis MRI" in paragraph_text,
    "draft_notice_present": "AI-assisted research draft—not submission-ready" in paragraph_text,
    "abstract_present": "Abstract" in paragraph_text,
    "conclusion_present": "8. Conclusion" in paragraph_text,
    "three_tables": len(document.tables) == 3,
    "one_figure": len(document.inline_shapes) == 1,
    "figure_alt_text": "MS MRI AI data-to-clinic evidence pathway" in document_xml,
    "reference_count_83": audit["reference_count"] == 83,
    "no_unresolved_citations": audit["unresolved_citations"] == [],
    "no_unused_references": audit["unused_references"] == [],
    "main_text_under_7000": audit["main_text_word_count_approx"] < 7000,
    "single_page_field": footer_xml.count("PAGE") == 1,
    "portrait_and_landscape_sections": {
        section.orientation for section in document.sections
    } == {WD_ORIENT.PORTRAIT, WD_ORIENT.LANDSCAPE},
    "no_source_citation_tokens": "[@" not in paragraph_text,
    "all_tables_titled": all(
        title in paragraph_text
        for title in (
            "Table 1. MS MRI datasets and benchmark characteristics",
            "Table 2. Task-specific technical performance and evidence limitations",
            "Table 3. Clinical translation readiness matrix",
        )
    ),
}

failed = [name for name, passed in checks.items() if not passed]
print(json.dumps({"checks": checks, "failed": failed}, indent=2))
if failed:
    raise SystemExit(1)
