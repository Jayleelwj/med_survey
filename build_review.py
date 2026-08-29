from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from references import REFERENCES


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "manuscript.md"
FIGURES = {
    "FIGURE1": {
        "path": ROOT / "output" / "figures" / "fig1_data_to_clinic_overview.png",
        "title": "MS MRI AI data-to-clinic evidence pathway",
        "alt": "A five-panel overview of MS MRI data, input integrity, AI analyses, evidence levels and clinician-verified use.",
    },
    "FIGURE2": {
        "path": ROOT / "output" / "figures" / "fig2_ms_mri_dataset_landscape.png",
        "title": "MS MRI dataset landscape",
        "alt": "Four panels compare open benchmarks, controlled cohorts, clinical realism and longitudinal depth, and data-integrity dimensions.",
    },
    "FIGURE3": {
        "path": ROOT / "output" / "figures" / "fig3_ms_mri_target_atlas.png",
        "title": "MRI targets for AI in multiple sclerosis",
        "alt": "Six MRI panels show lesion burden, longitudinal change, enhancement, cortical lesions, CVS and PRL, and spinal cord lesions.",
    },
    "FIGURE4": {
        "path": ROOT / "output" / "figures" / "fig4_evidence_translation_landscape.png",
        "title": "Evidence requirements for human-AI deployment",
        "alt": "Four panels show validation scope, generalisation under domain shift, decision evidence and monitored human-AI deployment.",
    },
}
OUTPUT = ROOT / "output" / "Artificial_Intelligence_in_Multiple_Sclerosis_MRI_Survey_Springer.docx"
AUDIT = ROOT / "output" / "springer_manuscript_audit.json"

INDIGO = "243B64"
TEAL = "2B7A78"
AMBER = "B77A16"
LIGHT_BLUE = "EAF0F7"
LIGHT_TEAL = "EAF5F3"
LIGHT_AMBER = "FAF2E4"
MID_GREY = "667085"
LIGHT_GREY = "F4F6F8"
GRID = "B9C2CC"


TABLE1 = {
    "title": "Table 1. MS MRI datasets and benchmark characteristics",
    "headers": [
        "Resource", "People / sessions", "Primary task and labels", "MRI", "Centres / scanners",
        "Longitudinal", "Annotation", "Access", "Principal limitation"
    ],
    "widths": [2.2, 1.8, 3.0, 2.5, 2.6, 1.9, 2.7, 2.0, 4.0],
    "rows": [
        ["ISBI 2015", "19 / 82", "WM lesion segmentation; hidden test masks", "T1, PD/T2, FLAIR", "1 site; 3 T Philips", "4–6 visits", "2 raters per session", "Registered public challenge", "Scan count obscures n=19; single scanner; high time-point leakage risk"],
        ["MSSEG 2016", "53 / 53", "Cross-sectional WM lesions", "3D FLAIR, T1±Gd, PD/T2", "3 sites; 4 scanners; 1.5/3 T", "No", "7 experts; LOP-STAPLE", "Controlled request; hidden test", "Only 15 labelled training cases; small unseen-scanner subset"],
        ["MSSEG-2", "100 / 200", "New-lesion segmentation between paired FLAIR", "3D FLAIR", "15 scanners; 1.5/3 T; 3 vendors", "2 visits, 1–3 y", "4 readers plus senior consensus", "Controlled challenge; hidden test", "FLAIR-only; vendor/site/label shifts entangled; 40% low-agreement lesions"],
        ["3D-MR-MS", "30 / 30", "Cross-sectional lesion masks", "T1, T2, FLAIR, post-Gd T1", "1 site; 3 T Siemens", "No", "3 experts; joint revision", "Public", "Small, single scanner; not all source contrasts native isotropic"],
        ["Long-MR-MS", "20 / 40", "New, disappearing, enlarging, shrinking lesions", "2D T1, T2, FLAIR", "1 site; 1.5 T Philips", "2 visits", "2 experts; consensus", "Public", "One interval per patient; 3-mm slices; labels depend on registration"],
        ["MSLesSeg", "75 / 115", "WM lesion segmentation", "T1, T2, FLAIR", "Several hospitals; ~3 1.5 T systems", "1–4 visits", "Two-reader workflow plus senior validation", "CC BY 4.0", "Most participants and all test cases have one time point"],
        ["MS-Baghdad", "60 / 60", "Lesion masks plus EDSS", "Routine T1, T2, FLAIR", "20 centres; 1.5 T", "No", "3 radiology/neurology experts", "Public", "Thick-slice heterogeneity; limited protocol metadata; mildly disabled sample"],
        ["SibBMS", "193 / 193", "MS/RIS and healthy controls; partial lesion labels", "T1, T2, FLAIR; post-Gd T1 in patients", "Single region; scanner incompletely reported", "No", "Primary description inconsistent about mask coverage", "Downloadable; licence wording unclear", "MS/RIS pooled; control-source shortcuts; annotation and licence ambiguity"],
        ["MS3SEG", "100 / 100", "MS lesions vs other WMH", "2D T1, T2, axial/sagittal FLAIR", "1 site; 1.5 T Toshiba/Canon", "No", "Junior–senior consensus; ≥3-mm lesions", "CC BY 4.0", "Single scanner; excludes sub-3-mm lesions; no outcomes"],
        ["PediMS", "9 / 28", "Paediatric lesions and clinical follow-up", "T1, T2, FLAIR", "One dominant 3 T Siemens site", "1–6 visits", "Junior rater plus specialist validation", "CC BY", "Extremely small; irregular intervals; best suited to OOD testing"],
        ["Shifts 2.0", "172 scan samples", "Uncertainty under distribution shift", "Common preprocessed T1/FLAIR", "Repackages public sets plus private Lausanne", "Mixed", "Source-dataset references", "Public parts; private evaluation server", "Not an independent cohort; compound population/scanner/label shift"],
        ["MS PATHS", "16,568 enrolled; 14,414 MRI studies", "Standardised quantitative MRI plus clinical outcomes", "3D MPRAGE + SPACE-FLAIR", "10 institutions; Siemens 3 T", "3,822 with repeat MRI", "No dense voxel masks at scale", "Controlled collaboration", "Vendor-specific standardisation; care-driven follow-up; industry involvement"],
        ["MindGlide trial ecosystem", "Train 2,871 / 4,247; external 1,001 / 14,952", "Tissue/lesion segmentation; longitudinal trial analyses", "Mixed legacy clinical/trial contrasts", "592 training scanners; multiple trials/cohorts", "Extensive", "Trial/cohort-derived labels", "Mostly proprietary", "Large and heterogeneous, but composition and labels limit independent reproduction"],
    ],
    "note": "People and imaging sessions are separated because repeat examinations are not independent. WM, white matter; Gd, gadolinium; RIS, radiologically isolated syndrome; OOD, out-of-distribution. Counts are tied to the cited descriptors and should not be summed across repackaged datasets."
}


TABLE2 = {
    "title": "Table 2. Task-specific technical performance and evidence limitations",
    "headers": ["Task", "Clinically relevant output", "Appropriate metrics", "Illustrative performance under stated conditions", "Highest common validation level", "Recurrent failure mode"],
    "widths": [3.0, 4.0, 4.0, 5.3, 3.4, 4.9],
    "rows": [
        ["Reconstruction / harmonisation", "Pathology-preserving image or stable measurement", "Downstream lesion/volume error; test–retest; site variance", "Super-resolution in simulated low-resolution FLAIR: PSNR 22.7→24.1 dB; SSIM 0.64→0.72", "Multi-dataset technical validation", "Hallucinated or smoothed small lesions; paired-data dependence; protocol drift"],
        ["Brain WM lesion segmentation", "Lesion map, count and volume", "Dice plus lesion sensitivity, PPV, FP/scan, volume bias", "Same-domain Dice commonly ~0.70–0.80; one real-world study 0.73 internal vs 0.66 external; inter-rater 0.63", "Small independent external sets", "Small/infratentorial lesions, low burden, vascular WMH, unfamiliar scanners"],
        ["New/enlarging lesions", "Patient activity and candidate changes", "Lesion F1; FP in stable cases; patient activity category", "MSSEG-2 best method and expert each ~85% for 0, 1–2, >2 new lesions; lesion overlap much more variable", "Hidden multi-scanner challenge; retrospective PACS", "Registration error, tiny lesions, protocol mismatch, false-positive review burden"],
        ["Gd-enhancing lesions", "Active enhancing lesion candidates", "Lesion sensitivity, FDR, FP/scan, size strata", "Representative Dice ~0.70–0.76; lesions >100 voxels materially easier", "Patient-level internal validation; routine retrospective comparison", "Rare positives, vessels, motion, tiny punctate enhancement"],
        ["Spinal cord and lesions", "Cord area and intramedullary lesion map", "Cord/lesion Dice; lesion precision and sensitivity; QC failure", "30-site cord Dice 0.95; MS lesion Dice 0.60, sensitivity 0.83, precision 0.77", "Multi-site internal/external technical testing", "Partial volume, motion/pulsation, small cord, missing axial or STIR/PSIR contrasts"],
        ["Cortical lesions", "Cortical lesion candidates by subtype", "Lesion F1/sensitivity, FP rate, subtype performance", "7 T detection ~74% with ~30% FP; multi-institution F1 ~0.64 in-domain vs ~0.50 OOD", "Very small external high-field cohorts", "Intracortical lesions, leukocortical boundary, 7 T/3 T domain shift"],
        ["CVS", "CVS-positive lesion proportion / rule", "Lesion AUC plus patient sensitivity, specificity, non-evaluable rate", "CVSnet AUC 0.90; 50% rule sensitivity 0.89, specificity 0.92 after manual lesion selection", "Small multi-scanner held-out test", "Manual candidate dependence; sequence and eligibility filtering"],
        ["PRL", "Rim-positive lesion count", "Sensitivity, specificity, PPV, FP/patient; reader agreement", "RimNet AUC 0.943, sensitivity 0.706, specificity 0.949, PPV 0.569; cross-validation dominant", "Internal cross-validation / limited multi-centre", "Extreme class imbalance, confluent/small lesions, specialised susceptibility imaging"],
        ["Diagnosis / differential", "Calibrated probability within a suspected-case spectrum", "Sensitivity, specificity, PPV/NPV, calibration, decision curve", "MS vs healthy controls may exceed 0.85 accuracy; MS vs NMOSD example AUC 0.85; pooled estimates are spectrum-sensitive", "Mostly same-ecosystem hold-out; few small external sets", "Healthy-control shortcuts, artificial prevalence, restricted mimics, missing clinical/CSF context"],
        ["Atrophy / tissue quantification", "Longitudinal volume change with uncertainty", "Bias/limits of agreement, test–retest, annualised change, QC rate", "High cross-sectional tissue Dice and correlations are common; annual change remains scanner/protocol sensitive", "Large trial validation; limited routine prospective use", "Scanner upgrade, segmentation drift, pseudo-label bias, threshold disagreement"],
        ["Disability / PIRA / cognition", "Calibrated future risk at a prespecified horizon", "AUC plus calibration, Brier score, RMSE, net benefit", "Systematic review: median AUC 0.78; pooled EDSS RMSE 1.31 with I²=95%", "Only a minority independently external", "Unstable outcome definitions, treatment confounding, small events, absent calibration"],
    ],
    "note": "Values are representative findings under specific datasets and validation levels, not pooled performance ranges. Dice and AUC must not be compared across tasks, lesion prevalences, label policies or test populations. PPV, positive predictive value; FP, false positive; FDR, false-discovery rate; CVS, central vein sign; PRL, paramagnetic rim lesion; PIRA, progression independent of relapse activity."
}


TABLE3 = {
    "title": "Table 3. Clinical translation readiness matrix",
    "headers": ["Application", "Technical validation", "Independent cross-centre evidence", "Prospective / workflow evidence", "Regulatory position", "Patient net-benefit evidence", "Readiness judgement"],
    "widths": [3.4, 3.2, 3.8, 3.9, 3.5, 3.5, 4.2],
    "rows": [
        ["Brain WM lesion segmentation", "High", "Moderate; external loss is consistent", "Limited reader studies", "Several cleared quantitative tools", "Absent", "Near clinical use as editable quantitative assistance"],
        ["Longitudinal new/enlarging lesion monitoring", "High for candidate detection", "Moderate; retrospective multi-centre signal", "Small crossover and time studies; failure burden documented", "Included in some cleared products under constrained inputs", "No completed outcome trial", "Strongest translational case as supervised second reader"],
        ["Brain/tissue atrophy quantification", "High cross-sectionally; moderate longitudinally", "Moderate in trials; protocol-sensitive", "Sparse", "Cleared volumetry tools exist", "Absent; cost models assumption-heavy", "Useful adjunct when acquisition and software versions are stable"],
        ["Cord contour / spinal lesion analysis", "High for cord; moderate for lesions", "Moderate technical evidence", "Absent", "No MS-specific patient-benefit indication identified", "Absent", "Cord morphometry nearer use than lesion automation"],
        ["Cortical lesions, CVS and PRL", "Moderate in specialist data", "Low to moderate; external sets small", "Absent", "No autonomous diagnostic indication identified", "Absent", "Research/specialist adjunct; end-to-end validation required"],
        ["MS diagnosis / differential diagnosis", "Moderate discrimination", "Low; spectrum often artificial", "Absent", "No autonomous MS diagnosis clearance identified", "Absent", "Technical validation stage; image-only use is inappropriate"],
        ["EDSS, PIRA, cognition or treatment prediction", "Moderate and heterogeneous", "Low; calibration rarely reported", "Absent", "No clinical prediction indication identified", "Absent", "Research stage; do not use for individual treatment choice"],
    ],
    "note": "Readiness integrates evidence quality rather than algorithm availability. Regulatory clearance is product-, version-, input- and intended-use-specific and is not equivalent to demonstrated workflow or patient benefit."
}


class CitationManager:
    def __init__(self) -> None:
        self.order: list[str] = []
        self.number: dict[str, int] = {}

    def register(self, key: str) -> int:
        if key not in REFERENCES:
            raise KeyError(f"Reference key not found: {key}")
        if key not in self.number:
            self.order.append(key)
            self.number[key] = len(self.order)
        return self.number[key]

    @staticmethod
    def compress(numbers: list[int]) -> str:
        nums = sorted(dict.fromkeys(numbers))
        parts: list[str] = []
        i = 0
        while i < len(nums):
            j = i
            while j + 1 < len(nums) and nums[j + 1] == nums[j] + 1:
                j += 1
            if j - i >= 2:
                parts.append(f"{nums[i]}–{nums[j]}")
            elif j - i == 1:
                parts.extend([str(nums[i]), str(nums[j])])
            else:
                parts.append(str(nums[i]))
            i = j + 1
        return ",".join(parts)

    def replace(self, text: str) -> str:
        def repl(match: re.Match[str]) -> str:
            keys = [x.strip().lstrip("@") for x in match.group(1).split(";")]
            numbers = [self.register(k) for k in keys]
            return f"[{self.compress(numbers)}]"
        return re.sub(r"\[@([^\]]+)\]", repl, text)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=45, start=60, bottom=45, end=60) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_number(paragraph) -> None:
    p = paragraph
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def set_section(section, landscape: bool = False) -> None:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)


def configure_page_numbers(doc: Document) -> None:
    for index, section in enumerate(doc.sections):
        section.footer.is_linked_to_previous = index > 0
        section.header.is_linked_to_previous = index > 0
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.clear()
    add_page_number(p)
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("Artificial intelligence in multiple sclerosis MRI · Survey Article")
    hr.font.name = "Times New Roman"
    hr.font.size = Pt(8)
    hr.font.italic = True
    hr.font.color.rgb = RGBColor.from_string(MID_GREY)


def add_inline(paragraph, text: str, cm: CitationManager, default_size=11, default_font="Times New Roman") -> None:
    text = cm.replace(text)
    token_re = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            run.font.name = default_font
            run.font.size = Pt(default_size)
        token = match.group(0)
        run = paragraph.add_run(token.strip("*"))
        run.font.name = default_font
        run.font.size = Pt(default_size)
        if token.startswith("**"):
            run.bold = True
        else:
            run.italic = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = default_font
        run.font.size = Pt(default_size)


def add_caption(doc: Document, text: str, cm: CitationManager, label_color=INDIGO) -> None:
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.keep_together = True
    add_inline(p, text, cm, default_size=8.5, default_font="Times New Roman")
    if p.runs:
        p.runs[0].font.color.rgb = RGBColor.from_string(label_color)


def add_figure(doc: Document, cm: CitationManager, spec: dict, legend: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(spec["path"]), width=Cm(15.8))
    inline = run._r.xpath(".//wp:inline")
    if inline:
        doc_pr = inline[0].find(qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("descr", spec["alt"])
            doc_pr.set("title", spec["title"])
    add_caption(doc, legend, cm)


def add_table(doc: Document, cm: CitationManager, spec: dict) -> None:
    dense_table = len(spec["headers"]) >= 8
    header_size = 6.8 if dense_table else 7.4
    body_size = 6.3 if dense_table else 7.0
    note_size = 6.5 if dense_table else 7.2
    cell_vertical_margin = 32 if dense_table else 55
    landscape = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section(landscape, landscape=True)
    cap = doc.add_paragraph()
    cap.paragraph_format.keep_with_next = True
    cap.paragraph_format.space_after = Pt(5)
    cap.paragraph_format.line_spacing = 1.0
    run = cap.add_run(spec["title"])
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0, 0, 0)

    table = doc.add_table(rows=1, cols=len(spec["headers"]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for idx, (cell, head, width) in enumerate(zip(table.rows[0].cells, spec["headers"], spec["widths"])):
        cell.width = Cm(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell, top=cell_vertical_margin, start=65, bottom=cell_vertical_margin, end=65)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(head)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(header_size)
        r.font.color.rgb = RGBColor(0, 0, 0)
    set_repeat_table_header(table.rows[0])
    set_keep_together(table.rows[0])

    for r_idx, data in enumerate(spec["rows"]):
        cells = table.add_row().cells
        set_keep_together(table.rows[-1])
        fill = "FFFFFF" if r_idx % 2 == 0 else LIGHT_GREY
        for cell, value, width in zip(cells, data, spec["widths"]):
            cell.width = Cm(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, fill)
            set_cell_margins(cell, top=cell_vertical_margin, start=65, bottom=cell_vertical_margin, end=65)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, value, cm, default_size=body_size, default_font="Times New Roman")
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.line_spacing = 1.0
    r = note.add_run("Note: ")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(note_size)
    add_inline(note, spec["note"], cm, default_size=note_size, default_font="Times New Roman")

    portrait = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section(portrait, landscape=False)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.widow_control = True
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title = styles["Title"]
    title.font.name = "Times New Roman"
    title.font.size = Pt(19)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.space_after = Pt(16)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for name, size, color in (("Heading 1", 14, INDIGO), ("Heading 2", 12, "1F2937"), ("Heading 3", 11.5, "1F2937")):
        st = styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        st.paragraph_format.space_after = Pt(5)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.line_spacing = 1.05

    caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(8.5)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(MID_GREY)

    if "Reference" not in styles:
        ref = styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref = styles["Reference"]
    ref.font.name = "Times New Roman"
    ref.font.size = Pt(9)
    ref.paragraph_format.line_spacing = 1.05
    ref.paragraph_format.space_after = Pt(3)
    ref.paragraph_format.left_indent = Cm(0.7)
    ref.paragraph_format.first_line_indent = Cm(-0.7)

    list_bullet = styles["List Bullet"]
    list_bullet.font.name = "Times New Roman"
    list_bullet.font.size = Pt(11)
    list_bullet.paragraph_format.line_spacing = 1.35
    list_bullet.paragraph_format.space_after = Pt(4)


def add_title_page(doc: Document, lines: list[str], cm: CitationManager) -> None:
    first = doc.sections[0]
    first.different_first_page_header_footer = True
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(16)
    for idx, line in enumerate(lines):
        if not line.strip():
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(16)
            p.paragraph_format.keep_with_next = True
            add_inline(p, line[2:], cm, default_size=19, default_font="Times New Roman")
            for run in p.runs:
                run.bold = True
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        add_inline(p, line, cm, default_size=10.5, default_font="Times New Roman")
        if line.startswith("**Article type:**"):
            p.paragraph_format.space_after = Pt(12)
        if line.startswith("**Draft status:**"):
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(10)
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor.from_string(MID_GREY)
            set_cell_like_paragraph(p, LIGHT_GREY)
    doc.add_page_break()


def set_cell_like_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def build() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    main_text_source = text[text.index("## 1. Introduction"):text.index("## Figure legends")]
    main_text_words = len(re.findall(r"\b[\w’'-]+\b", re.sub(r"\[@[^\]]+\]", "", main_text_source)))
    text = text.replace("[Insert final word count]", f"{main_text_words:,}")
    lines = text.splitlines()
    figure_heading_index = lines.index("## Figure legends")
    declarations_index = lines.index("## Declarations")
    legends: dict[str, str] = {}
    for legend_line in lines[figure_heading_index + 1:declarations_index]:
        match = re.match(r"\*\*Figure (\d+)\.", legend_line.strip())
        if match:
            visible_legend = re.sub(r"\s+\*\*Alt text:\*\*.*$", "", legend_line.strip())
            legends[f"FIGURE{match.group(1)}"] = visible_legend
    missing_legends = sorted(set(FIGURES) - set(legends))
    if missing_legends:
        raise RuntimeError(f"Missing figure legends: {missing_legends}")
    # Omit the separate legend section because each legend is placed with its inline figure.
    lines = lines[:figure_heading_index] + lines[declarations_index:]

    abstract_index = lines.index("## Abstract")
    title_lines = lines[:abstract_index]
    body_lines = lines[abstract_index:]

    doc = Document()
    configure_styles(doc)
    set_section(doc.sections[0], landscape=False)
    cm = CitationManager()
    add_title_page(doc, title_lines, cm)

    i = 0
    while i < len(body_lines):
        raw = body_lines[i]
        line = raw.strip()
        i += 1
        if not line:
            continue
        figure_match = re.fullmatch(r"\[\[(FIGURE\d+)\]\]", line)
        if figure_match:
            figure_key = figure_match.group(1)
            if figure_key not in FIGURES:
                raise RuntimeError(f"Unknown figure marker: {figure_key}")
            add_figure(doc, cm, FIGURES[figure_key], legends[figure_key])
            continue
        if line == "[[TABLE1]]":
            add_table(doc, cm, TABLE1)
            continue
        if line == "[[TABLE2]]":
            add_table(doc, cm, TABLE2)
            continue
        if line == "[[TABLE3]]":
            add_table(doc, cm, TABLE3)
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:], cm, default_size=12, default_font="Times New Roman")
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, line[3:], cm, default_size=14, default_font="Times New Roman")
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:], cm)
            continue
        p = doc.add_paragraph(style="Normal")
        add_inline(p, line, cm)
        if line.startswith("**Keywords:**"):
            p.paragraph_format.space_after = Pt(6)

    unresolved = sorted(set(re.findall(r"@([A-Za-z0-9_]+)", text)) - set(REFERENCES))
    if unresolved:
        raise RuntimeError(f"Unresolved citation keys: {unresolved}")
    unused = sorted(set(REFERENCES) - set(cm.order))
    if unused:
        raise RuntimeError(f"Registry entries not cited: {unused}")

    doc.add_page_break()
    h = doc.add_paragraph(style="Heading 1")
    h.add_run("References")
    for n, key in enumerate(cm.order, start=1):
        p = doc.add_paragraph(style="Reference")
        p.add_run(f"{n}. ").bold = True
        p.add_run(REFERENCES[key])

    # Core properties identify the document as a research draft.
    doc.core_properties.title = "A Survey of Artificial Intelligence in Multiple Sclerosis MRI"
    doc.core_properties.subject = "Datasets, technical performance, evidence quality, and clinical translation"
    doc.core_properties.author = "[Author placeholders]"
    doc.core_properties.keywords = "multiple sclerosis; MRI; artificial intelligence; datasets; validation; clinical translation"
    doc.core_properties.comments = "AI-assisted research draft—not submission-ready"

    configure_page_numbers(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)

    audit = {
        "document": str(OUTPUT),
        "main_text_word_count_approx": main_text_words,
        "total_source_word_count_approx": len(re.findall(r"\b[\w’'-]+\b", re.sub(r"\[@[^\]]+\]", "", text))),
        "reference_count": len(cm.order),
        "reference_order": cm.order,
        "figures": [str(FIGURES[key]["path"]) for key in sorted(FIGURES)],
        "tables": [TABLE1["title"], TABLE2["title"], TABLE3["title"]],
        "unresolved_citations": unresolved,
        "unused_references": unused,
    }
    AUDIT.write_text(json.dumps(audit, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps(audit, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    build()
